from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from markdown_formatter import DocumentBlock, parse_markdown_blocks


DEFAULT_FONT = "Vazirmatn"


def _set_run_font(run, font_name: str, size: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts

    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)

    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)


def _set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.35

    paragraph_properties = paragraph._p.get_or_add_pPr()

    bidi = paragraph_properties.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        paragraph_properties.append(bidi)

    bidi.set(qn("w:val"), "1")


def _set_document_defaults(document: Document, font_name: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = document.styles

    normal_style = styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(13)

    for style_name, size in (("Title", 20), ("Heading 1", 17), ("Heading 2", 15), ("Heading 3", 14)):
        style = styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True


def _add_block(document: Document, block: DocumentBlock, font_name: str) -> None:
    if block.block_type == "separator":
        paragraph = document.add_paragraph()
        _set_rtl(paragraph)
        run = paragraph.add_run("ــــــــــــــــــــــــــــــــــــــــ")
        _set_run_font(run, font_name, 10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if block.block_type == "heading":
        paragraph = document.add_paragraph()
        _set_rtl(paragraph)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)

        size = {1: 17, 2: 15, 3: 14}.get(block.level, 14)
        run = paragraph.add_run(block.text)
        _set_run_font(run, font_name, size, bold=True)
        return

    if block.block_type in {"bullet", "numbered"}:
        paragraph = document.add_paragraph()
        _set_rtl(paragraph)
        paragraph.paragraph_format.right_indent = Cm(0.45)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)

        prefix = "• " if block.block_type == "bullet" else ""
        run = paragraph.add_run(f"{prefix}{block.text}")
        _set_run_font(run, font_name, 13)
        return

    if block.block_type == "quote":
        paragraph = document.add_paragraph()
        _set_rtl(paragraph)
        paragraph.paragraph_format.right_indent = Cm(0.7)
        paragraph.paragraph_format.left_indent = Cm(0.7)

        run = paragraph.add_run(f"«{block.text}»")
        _set_run_font(run, font_name, 12.5)
        run.italic = True
        return

    paragraph = document.add_paragraph()
    _set_rtl(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0.65)

    run = paragraph.add_run(block.text)
    _set_run_font(run, font_name, 13)


def create_word_document(
    title: str,
    markdown_content: str,
    font_name: str = DEFAULT_FONT,
) -> bytes:
    document = Document()
    _set_document_defaults(document, font_name)

    clean_title = str(title or "").strip() or "مقاله پژوهشی"

    title_paragraph = document.add_paragraph()
    _set_rtl(title_paragraph)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(18)

    title_run = title_paragraph.add_run(clean_title)
    _set_run_font(title_run, font_name, 20, bold=True)

    subtitle = document.add_paragraph()
    _set_rtl(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    subtitle_run = subtitle.add_run(
        "برگرفته از پایگاه جامع درسی استاد علامه سید علی موسوی(ره)"
    )
    _set_run_font(subtitle_run, font_name, 11)

    blocks: Iterable[DocumentBlock] = parse_markdown_blocks(markdown_content)

    for block in blocks:
        _add_block(document, block, font_name)

    footer = document.sections[0].footer.paragraphs[0]
    _set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_run = footer.add_run("تولیدشده توسط سامانه ProfessorAI")
    _set_run_font(footer_run, font_name, 9)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
